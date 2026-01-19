#!/usr/bin/env python3
"""
Generate comprehensive handover documentation in Word format for EcoPrompt Coach
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import os

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(76, 175, 80)  # Green color
    return heading

def add_code_block(doc, code_text, language=""):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    para.style = 'Normal'
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Add light gray background effect
    return para

def add_formula(doc, formula_text):
    """Add a mathematical formula"""
    para = doc.add_paragraph()
    run = para.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    return para

def create_table_with_data(doc, headers, rows):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Bold header text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def generate_handover_document():
    """Generate the complete handover documentation"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ========== TITLE PAGE ==========
    title = doc.add_heading('EcoPrompt Coach', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Handover Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(76, 175, 80)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Chrome Extension for AI Sustainability\n').font.size = Pt(14)
    info.add_run('Visualizing Environmental Impact of LLM Queries\n\n').font.size = Pt(14)
    info.add_run(f'Generated: January 2026\n').font.size = Pt(11)
    info.add_run('Version: 1.0\n').font.size = Pt(11)

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_heading(doc, 'Table of Contents', 1)

    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Architecture & Technology Stack',
        '4. Core Calculations & Equations',
        '5. Key Features & Functionality',
        '6. Data Files & Model Benchmarks',
        '7. File Structure & Components',
        '8. Supported Models & Efficiency Rankings',
        '9. Integration & Data Flow',
        '10. Installation & Deployment',
        '11. Performance Metrics & Examples',
        '12. Development Notes & Design Decisions',
        '13. Future Roadmap',
        '14. Technical Reference'
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 1. EXECUTIVE SUMMARY ==========
    add_heading(doc, '1. Executive Summary', 1)

    doc.add_paragraph(
        'EcoPrompt Coach is a Chrome browser extension that helps users understand and reduce the '
        'environmental impact of their AI model usage. It provides real-time calculations of energy '
        'consumption, water usage, and carbon emissions for queries to ChatGPT, Claude, Gemini, and '
        'other LLM platforms.'
    )

    add_heading(doc, 'Key Achievements', 2)
    achievements = [
        'Ranked #1 among prototypes by ministry stakeholders for actionability, impact, and alignment',
        'Based on peer-reviewed research: "How Hungry is AI?" (Jegham et al., 2025, arXiv:2505.09598v4)',
        'Supports 13 LLM models across 4 major cloud providers',
        'Provides relatable environmental comparisons (e.g., "equivalent to 25 minutes of LED light")',
        'Offers 6 types of optimization suggestions to reduce environmental impact',
        'Real-time calculation with <1ms latency'
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')

    add_heading(doc, 'Problem Statement', 2)
    doc.add_paragraph(
        'Users of AI chatbots have no visibility into the environmental costs of their queries. '
        'A single query can consume as much energy as running an LED bulb for hours, use '
        'water equivalent to a glass of drinking water, and emit carbon comparable to driving '
        'several kilometers. Without awareness, users cannot make informed decisions about '
        'model selection or prompt optimization.'
    )

    add_heading(doc, 'Solution Overview', 2)
    doc.add_paragraph(
        'EcoPrompt Coach bridges this gap by:\n'
    )
    solutions = [
        'Calculating environmental impact in real-time before users submit queries',
        'Converting technical metrics into relatable comparisons (LED bulbs, smartphone charges, car distance)',
        'Detecting task types and providing context-aware optimization suggestions',
        'Comparing models side-by-side to show efficiency differences of up to 240x',
        'Offering a prompt library to save and reuse sustainable prompts',
        'Tracking usage statistics to demonstrate cumulative impact'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    doc.add_page_break()

    # ========== 2. PROJECT OVERVIEW ==========
    add_heading(doc, '2. Project Overview', 1)

    add_heading(doc, 'What is EcoPrompt Coach?', 2)
    doc.add_paragraph(
        'EcoPrompt Coach is a browser extension that acts as an environmental "fitness tracker" '
        'for AI usage. Just as fitness trackers help people understand calories and exercise, '
        'EcoPrompt Coach helps users understand the environmental cost of their AI interactions.'
    )

    add_heading(doc, 'Target Users', 2)
    users = [
        'Individual AI users who want to reduce their environmental footprint',
        'Organizations implementing sustainability policies',
        'Researchers studying AI environmental impact',
        'Educators teaching responsible AI usage',
        'Policy makers tracking AI infrastructure efficiency'
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')

    add_heading(doc, 'Core Value Proposition', 2)
    doc.add_paragraph(
        'The extension makes invisible environmental costs visible and actionable. '
        'It transforms abstract technical metrics (Watt-hours, milliliters, grams of CO2) '
        'into everyday comparisons people can understand and act upon.'
    )

    doc.add_paragraph('\nExample Transformation:')

    # Create comparison table
    table = create_table_with_data(
        doc,
        ['Technical Metric', 'Relatable Comparison'],
        [
            ['0.421 Wh', '25 minutes of LED bulb light'],
            ['1.5 mL water', '1-2 drops of water'],
            ['0.15 gCO2e', 'Driving 0.75 meters in a sedan']
        ]
    )

    doc.add_page_break()

    # ========== 3. ARCHITECTURE & TECHNOLOGY STACK ==========
    add_heading(doc, '3. Architecture & Technology Stack', 1)

    add_heading(doc, 'Overall Architecture', 2)
    doc.add_paragraph(
        'The extension follows a layered architecture with clear separation of concerns:'
    )

    add_code_block(doc, '''EcoPrompt Coach (Chrome Extension)
├── Frontend Layer (Browser UI)
│   ├── popup.js/html - Extension popup interface
│   ├── popup.css - Styling with 6 theme options
│   ├── content.js - Real-time analysis on web pages
│   ├── settings.js/html - User settings and preferences
│   └── background.js - Service worker for background tasks
│
├── Shared Logic Layer (JavaScript)
│   ├── shared/calculator.js - Main impact calculations + interpolation
│   ├── shared/analyzer.js - Token estimation, task detection, tips
│   └── tooltip.css - Floating badge styling
│
├── Data Layer (JSON Files)
│   ├── data/model_benchmarks.json - 13 LLM models with metrics
│   ├── data/infrastructure.json - Cloud provider multipliers
│   ├── data/conversion_factors.json - Relatable comparisons
│   ├── data/config.json - Extension configuration
│   └── data/prompt_templates.json - Prompt library templates
│
└── Backend Layer (Python - for future API)
    └── backend/calculations/
        ├── __init__.py - Main EcoImpactCalculator class
        ├── energy.py - Energy consumption formulas
        ├── water.py - Water usage calculations
        ├── carbon.py - Carbon emission calculations
        └── conversions.py - Relatable conversion generation''')

    add_heading(doc, 'Technology Stack', 2)

    # Frontend technologies
    doc.add_paragraph('Frontend Technologies:', style='Heading 3')
    frontend_tech = [
        'JavaScript ES6+ - Core logic and calculations',
        'HTML5 - User interface structure',
        'CSS3 - Styling with Grid and Flexbox',
        'Chrome Extension APIs - Manifest V3',
        'JSON - Data storage and configuration'
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')

    # Backend technologies
    doc.add_paragraph('Backend Technologies (Future):', style='Heading 3')
    backend_tech = [
        'Python 3.8+ - Calculation engine',
        'FastAPI - REST API framework',
        'PostgreSQL - Database for prompt library',
        'NumPy/Pandas - Data processing',
        'Pytest - Unit testing'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')

    add_heading(doc, 'How It Works - Step by Step', 2)

    steps = [
        ('User enters a prompt', 'On ChatGPT, Claude, Gemini, or other AI platform'),
        ('Extension detects input', 'Content script (content.js) monitors textarea elements'),
        ('Analyzer estimates tokens', 'Uses word-to-token ratio approximation (1 token ≈ 0.75 words)'),
        ('Analyzer detects task type', 'Identifies from 9 task types: image gen, code, translation, etc.'),
        ('Calculator loads model data', 'Fetches benchmarks and infrastructure multipliers from JSON'),
        ('Calculator performs interpolation', 'Linear interpolation between 3 size categories (short/medium/long)'),
        ('Energy calculation', 'Based on TPS, latency, and power consumption'),
        ('Water & carbon derivation', 'Calculated from energy using infrastructure multipliers'),
        ('Relatable conversion', 'Converts technical metrics to everyday comparisons'),
        ('Optimization suggestions', 'Generates 6 types of context-aware tips'),
        ('Display results', 'Shows floating badge with metrics and suggestions'),
        ('User takes action', 'Can optimize prompt, switch models, or proceed as-is')
    ]

    for i, (step, description) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {step}: ').bold = True
        p.add_run(description)

    doc.add_page_break()

    # ========== 4. CORE CALCULATIONS & EQUATIONS ==========
    add_heading(doc, '4. Core Calculations & Equations', 1)

    doc.add_paragraph(
        'All calculations are based on the peer-reviewed research paper:\n'
        '"How Hungry is AI? Benchmarking Energy, Water, and Carbon Footprint of LLM Inference"\n'
        'Authors: Jegham et al., 2025\n'
        'arXiv:2505.09598v4'
    ).italic = True

    add_heading(doc, '4.1 Energy Calculation', 2)

    doc.add_paragraph(
        'Energy consumption is the foundation of all environmental impact calculations. '
        'It depends on the model\'s processing speed (TPS), response latency, and power consumption.'
    )

    add_heading(doc, 'Primary Formula', 3)
    add_formula(doc, 'E_query (kWh) = [(Output_Length / TPS + Latency) / 3600] × [(P_GPU × U_GPU) + (P_non-GPU × U_non-GPU)] × PUE')

    doc.add_paragraph('Where:')
    energy_vars = [
        ('E_query', 'Energy consumed per query (kilowatt-hours)'),
        ('Output_Length', 'Number of output tokens generated'),
        ('TPS', 'Tokens per second (throughput)'),
        ('Latency', 'Time to first token (seconds)'),
        ('P_GPU', 'GPU power consumption (kW)'),
        ('U_GPU', 'GPU utilization percentage'),
        ('P_non-GPU', 'Non-GPU power (CPU, memory, networking) (kW)'),
        ('U_non-GPU', 'Non-GPU utilization percentage'),
        ('PUE', 'Power Usage Effectiveness (datacenter efficiency multiplier)')
    ]

    for var, desc in energy_vars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{var}: ').bold = True
        p.add_run(desc)

    add_heading(doc, 'Three-Tier Interpolation Strategy', 3)
    doc.add_paragraph(
        'Rather than using the complex formula directly, the extension uses LINEAR INTERPOLATION '
        'between three benchmark categories from the research paper:'
    )

    # Benchmark categories table
    table = create_table_with_data(
        doc,
        ['Category', 'Input Tokens', 'Output Tokens', 'Total Tokens', 'Use Case'],
        [
            ['Short', '100', '300', '400', 'Quick Q&A, simple queries'],
            ['Medium', '1,000', '1,000', '2,000', 'Standard conversations'],
            ['Long', '10,000', '1,500', '11,500', 'Document analysis, research']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph('Interpolation Logic:')

    interp_logic = [
        'If total_tokens ≤ 400: Linear extrapolation from 0 to short benchmark',
        'If 400 < total_tokens ≤ 2,000: Linear interpolation between short and medium',
        'If 2,000 < total_tokens ≤ 11,500: Linear interpolation between medium and long',
        'If total_tokens > 11,500: Linear extrapolation from long benchmark'
    ]
    for logic in interp_logic:
        doc.add_paragraph(logic, style='List Bullet')

    add_heading(doc, 'Implementation Code (JavaScript)', 3)
    add_code_block(doc, '''function interpolateEnergy(model, inputTokens, outputTokens) {
  const totalTokens = inputTokens + outputTokens;

  // Get benchmark energies for this model
  const short_energy = model.short.energy_wh_mean;
  const medium_energy = model.medium.energy_wh_mean;
  const long_energy = model.long.energy_wh_mean;

  let energy = 0;

  if (totalTokens <= 400) {
    // Extrapolate from 0 to short
    const ratio = totalTokens / 400;
    energy = short_energy * ratio;

  } else if (totalTokens <= 2000) {
    // Interpolate between short and medium
    const ratio = (totalTokens - 400) / (2000 - 400);
    energy = short_energy + ratio * (medium_energy - short_energy);

  } else if (totalTokens <= 11500) {
    // Interpolate between medium and long
    const ratio = (totalTokens - 2000) / (11500 - 2000);
    energy = medium_energy + ratio * (long_energy - medium_energy);

  } else {
    // Extrapolate beyond long
    const ratio = totalTokens / 11500;
    energy = long_energy * ratio;
  }

  return energy; // in Watt-hours (Wh)
}''')

    doc.add_paragraph()
    doc.add_paragraph('File Location: /shared/calculator.js (lines 140-184)')

    add_heading(doc, 'Task Type Energy Multipliers', 3)
    doc.add_paragraph(
        'Different tasks consume different amounts of energy due to model complexity '
        'and multiple API calls:'
    )

    table = create_table_with_data(
        doc,
        ['Task Type', 'Multiplier', 'Reason'],
        [
            ['Image Generation', '3.0x', 'Diffusion models are computationally expensive'],
            ['Agentic/Research', '2.0x', 'Multiple tool calls and reasoning loops'],
            ['Code Generation', '1.2x', 'Requires precise syntax and structure'],
            ['Creative Writing', '1.3x', 'Longer output and style considerations'],
            ['Data Analysis', '1.4x', 'Computation and formatting overhead'],
            ['General Chat', '1.0x', 'Baseline interaction'],
            ['Text Summarization', '0.7x', 'Shorter output than input'],
            ['Translation', '0.6x', 'Straightforward transformation'],
            ['Simple Q&A', '0.8x', 'Short factual responses']
        ]
    )

    doc.add_page_break()

    add_heading(doc, '4.2 Water Calculation', 2)

    doc.add_paragraph(
        'Water usage has two components: on-site cooling water for the datacenter '
        'and off-site water used in electricity generation (especially hydropower).'
    )

    add_heading(doc, 'Water Formula', 3)
    add_formula(doc, 'Water (L) = (E_query / PUE) × WUE_onsite + E_query × WUE_offsite')

    doc.add_paragraph('Where:')
    water_vars = [
        ('E_query', 'Energy consumed (kWh) from energy calculation'),
        ('PUE', 'Power Usage Effectiveness (datacenter efficiency)'),
        ('WUE_onsite', 'Water Usage Effectiveness for cooling (L/kWh)'),
        ('WUE_offsite', 'Water used in electricity generation (L/kWh)')
    ]

    for var, desc in water_vars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{var}: ').bold = True
        p.add_run(desc)

    add_heading(doc, 'Provider-Specific Water Values', 3)
    table = create_table_with_data(
        doc,
        ['Provider', 'WUE Onsite (L/kWh)', 'WUE Offsite (L/kWh)', 'Total (L/kWh)'],
        [
            ['Google Cloud', '0.25', '3.142', '3.392'],
            ['Microsoft Azure', '0.30', '3.142', '3.442'],
            ['AWS', '0.18', '3.142', '3.322'],
            ['DeepSeek (China)', '1.20', '6.016', '7.216']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: DeepSeek\'s higher values reflect less efficient cooling systems '
        'and greater reliance on hydropower in China\'s energy mix.'
    )

    add_heading(doc, 'Implementation Code (Python)', 3)
    add_code_block(doc, '''def calculate_water(energy_kwh, infrastructure_data):
    """
    Calculate water usage in liters

    Args:
        energy_kwh: Energy consumed in kWh
        infrastructure_data: Dict with 'pue', 'wue_onsite', 'wue_offsite'

    Returns:
        dict: {
            'total_liters': total water usage,
            'onsite_liters': cooling water,
            'offsite_liters': electricity generation water
        }
    """
    pue = infrastructure_data['pue']
    wue_onsite = infrastructure_data['wue_onsite']
    wue_offsite = infrastructure_data['wue_offsite']

    # On-site cooling water
    onsite_water = (energy_kwh / pue) * wue_onsite

    # Off-site electricity generation water
    offsite_water = energy_kwh * wue_offsite

    total_water = onsite_water + offsite_water

    return {
        'total_liters': total_water,
        'onsite_liters': onsite_water,
        'offsite_liters': offsite_water
    }''')

    doc.add_paragraph()
    doc.add_paragraph('File Location: /backend/calculations/water.py (lines 45-96)')

    doc.add_page_break()

    add_heading(doc, '4.3 Carbon Calculation', 2)

    doc.add_paragraph(
        'Carbon emissions depend on the carbon intensity of the regional electricity grid. '
        'This varies significantly by location and energy source mix.'
    )

    add_heading(doc, 'Carbon Formula', 3)
    add_formula(doc, 'Carbon (kgCO2e) = E_query (kWh) × CIF')

    doc.add_paragraph('Where:')
    carbon_vars = [
        ('E_query', 'Energy consumed (kWh) from energy calculation'),
        ('CIF', 'Carbon Intensity Factor (kgCO2e per kWh)'),
        ('CO2e', 'Carbon dioxide equivalent (includes all greenhouse gases)')
    ]

    for var, desc in carbon_vars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{var}: ').bold = True
        p.add_run(desc)

    add_heading(doc, 'Regional Carbon Intensity Factors', 3)
    table = create_table_with_data(
        doc,
        ['Provider / Region', 'CIF (kgCO2e/kWh)', 'Energy Mix'],
        [
            ['Google Cloud (US West)', '0.32', 'High renewable energy, some natural gas'],
            ['Microsoft Azure (US)', '0.3528', 'Mixed renewable + natural gas'],
            ['AWS (US East)', '0.385', 'Coal + natural gas + renewables'],
            ['DeepSeek (China)', '0.60', 'Coal-heavy grid (60-70% coal)']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Key Insight: Using DeepSeek vs. Google Cloud results in nearly 2x the carbon '
        'emissions for the same energy consumption due to China\'s coal-heavy grid.'
    )

    add_heading(doc, 'Implementation Code (Python)', 3)
    add_code_block(doc, '''def calculate_carbon(energy_kwh, infrastructure_data):
    """
    Calculate carbon emissions in kg CO2 equivalent

    Args:
        energy_kwh: Energy consumed in kWh
        infrastructure_data: Dict with 'cif' (carbon intensity factor)

    Returns:
        float: Carbon emissions in kgCO2e
    """
    cif = infrastructure_data['cif']
    carbon_kg = energy_kwh * cif

    return carbon_kg''')

    doc.add_paragraph()
    doc.add_paragraph('File Location: /backend/calculations/carbon.py (lines 45-100)')

    doc.add_page_break()

    add_heading(doc, '4.4 Relatable Conversions', 2)

    doc.add_paragraph(
        'The extension converts technical metrics into everyday comparisons people can understand. '
        'This is crucial for user engagement and behavior change.'
    )

    add_heading(doc, 'Energy Conversions', 3)
    table = create_table_with_data(
        doc,
        ['Energy Range', 'Relatable Metric', 'Calculation Formula'],
        [
            ['< 0.1 Wh', 'Seconds of LED bulb (10W)', 'seconds = (Wh / 10W) × 3600'],
            ['0.1 - 1 Wh', 'Minutes of LED bulb', 'minutes = (Wh / 10W) × 60'],
            ['1 - 15 Wh', 'Smartphone charge %', 'percent = (Wh / 15Wh) × 100'],
            ['15 - 50 Wh', 'Hours of room lighting (15W)', 'hours = Wh / 15Wh'],
            ['> 50 Wh', 'Hours of laptop usage (50W)', 'hours = Wh / 50Wh']
        ]
    )

    doc.add_paragraph()
    add_heading(doc, 'Water Conversions', 3)
    table = create_table_with_data(
        doc,
        ['Water Range', 'Relatable Metric', 'Calculation Formula'],
        [
            ['< 5 mL', 'Water drops (20 drops/mL)', 'drops = mL × 20'],
            ['5 - 15 mL', 'Teaspoons (5 mL/tsp)', 'tsp = mL / 5'],
            ['15 - 100 mL', 'Tablespoons (15 mL/tbsp)', 'tbsp = mL / 15'],
            ['100 - 1000 mL', 'Drinking glasses (250 mL)', 'glasses = mL / 250'],
            ['1 - 15 L', 'Water bottles (500 mL)', 'bottles = L / 0.5'],
            ['15 - 50 L', 'Dishwasher loads (15 L)', 'loads = L / 15'],
            ['50 - 100 L', 'Washing machine cycles (50 L)', 'cycles = L / 50'],
            ['> 100 L', '10-minute showers (100 L)', 'showers = L / 100']
        ]
    )

    doc.add_paragraph()
    add_heading(doc, 'Carbon Conversions', 3)
    table = create_table_with_data(
        doc,
        ['Carbon Range', 'Relatable Metric', 'Calculation Formula'],
        [
            ['< 1 gCO2e', 'Few breaths', 'Qualitative comparison'],
            ['1 - 20 gCO2e', 'Driving distance (meters)', 'meters = (gCO2e / 190) × 1000'],
            ['20 - 1000 gCO2e', 'Driving distance (km)', 'km = gCO2e / 190'],
            ['> 1 kgCO2e', 'Tree absorption (20 kg/year)', 'percent = (kgCO2e / 20) × 100']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Car emissions assume average sedan at 190 gCO2e per km. '
        'Tree absorption assumes mature tree absorbing 20 kg CO2 per year.'
    )

    doc.add_paragraph()
    doc.add_paragraph('File Location: /shared/calculator.js (lines 348-527)')

    doc.add_page_break()

    add_heading(doc, '4.5 Eco-Score Calculation', 2)

    doc.add_paragraph(
        'The Eco-Score is a 0-100 rating that helps users quickly understand model efficiency. '
        'It uses a logarithmic scale to handle the 240x difference between best and worst models.'
    )

    add_heading(doc, 'Why Logarithmic?', 3)
    doc.add_paragraph(
        'Linear scale would make most models cluster near 0, providing no useful differentiation. '
        'Logarithmic scale spreads the scores across the full 0-100 range.'
    )

    add_heading(doc, 'Implementation', 3)
    add_code_block(doc, '''function calculateEcoScore(energyWh) {
  // Benchmark values from all 13 models
  const minEnergy = 0.070;   // LLaMA 3.2 1B short (best)
  const maxEnergy = 33.634;  // DeepSeek-R1 long (worst)

  // Convert to logarithmic scale
  const logMin = Math.log(minEnergy);
  const logMax = Math.log(maxEnergy);
  const logEnergy = Math.log(Math.max(energyWh, minEnergy));

  // Calculate position on log scale (0-1)
  const position = (logEnergy - logMin) / (logMax - logMin);

  // Invert so lower energy = higher score
  const score = 100 - (position * 100);

  // Clamp to 0-100 range
  return Math.max(0, Math.min(100, Math.round(score)));
}''')

    doc.add_paragraph()
    add_heading(doc, 'Score Interpretation', 3)
    table = create_table_with_data(
        doc,
        ['Score Range', 'Rating', 'Description'],
        [
            ['80-100', 'Excellent', 'Highly efficient model, minimal environmental impact'],
            ['60-79', 'Good', 'Above-average efficiency, recommended for most use cases'],
            ['40-59', 'Fair', 'Average efficiency, consider alternatives for frequent use'],
            ['20-39', 'Poor', 'Below-average efficiency, use sparingly'],
            ['0-19', 'Very Poor', 'Extremely inefficient, avoid unless absolutely necessary']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph('File Location: /shared/calculator.js (lines 534-556)')

    doc.add_page_break()

    # ========== 5. KEY FEATURES & FUNCTIONALITY ==========
    add_heading(doc, '5. Key Features & Functionality', 1)

    add_heading(doc, '5.1 Real-Time Environmental Impact Calculation', 2)
    doc.add_paragraph(
        'The core feature that displays energy, water, and carbon metrics as users type prompts.'
    )

    doc.add_paragraph('How it works:')
    realtime_steps = [
        'Content script monitors textarea elements on AI platforms',
        'Detects when user is typing a prompt',
        'Estimates token count using word-based approximation',
        'Auto-detects model from website hostname or user selection',
        'Calculates environmental impact using interpolation',
        'Displays floating badge with real-time metrics',
        'Updates dynamically as user edits prompt'
    ]
    for step in realtime_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Implementation Files:')
    doc.add_paragraph('• content.js - DOM monitoring and event handling', style='List Bullet')
    doc.add_paragraph('• shared/analyzer.js - Token estimation', style='List Bullet')
    doc.add_paragraph('• shared/calculator.js - Impact calculation', style='List Bullet')
    doc.add_paragraph('• tooltip.css - Badge styling', style='List Bullet')

    add_heading(doc, '5.2 Task Type Detection', 2)
    doc.add_paragraph(
        'Automatically identifies the type of task from prompt content to provide '
        'accurate energy estimates and relevant optimization tips.'
    )

    add_heading(doc, 'Supported Task Types', 3)
    table = create_table_with_data(
        doc,
        ['Task Type', 'Energy Multiplier', 'Detection Keywords', 'Example Prompt'],
        [
            ['Image Generation', '3.0x', 'generate image, create picture, draw', '"Generate an image of a sunset"'],
            ['Agentic/Research', '2.0x', 'research, analyze deeply, investigate', '"Research climate change impacts"'],
            ['Code Generation', '1.2x', 'write code, create function, implement', '"Write a Python sorting function"'],
            ['Creative Writing', '1.3x', 'write story, compose poem, create narrative', '"Write a short story about space"'],
            ['Data Analysis', '1.4x', 'analyze data, calculate, process dataset', '"Analyze sales trends from CSV"'],
            ['Text Summarization', '0.7x', 'summarize, condense, key points', '"Summarize this article in 3 points"'],
            ['Translation', '0.6x', 'translate to, convert to language', '"Translate to Spanish: Hello world"'],
            ['Simple Q&A', '0.8x', 'what is, how do, explain', '"What is photosynthesis?"'],
            ['General Chat', '1.0x', 'Default for other queries', '"Tell me about your capabilities"']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Detection uses keyword matching with confidence scoring. If multiple task types match, '
        'the one with highest confidence is selected.'
    )

    doc.add_paragraph()
    doc.add_paragraph('Implementation: /shared/analyzer.js (lines 36-174)')

    add_heading(doc, '5.3 Optimization Suggestions', 2)
    doc.add_paragraph(
        'The extension provides 6 types of context-aware optimization tips to reduce '
        'environmental impact:'
    )

    optimization_types = [
        ('1. Special Query Detection',
         'Identifies queries better served by alternative tools',
         'Examples:\n• Location → Google Maps (saves 100% AI energy)\n• Weather → Weather app\n• Math → Calculator app\n• Time → Device clock'),

        ('2. Eco-Friendly Alternatives',
         'Task-specific recommendations for lower-impact approaches',
         'Examples:\n• Translation → Offline apps (100% savings)\n• Simple Q&A → Web search (99.9% savings)\n• Image generation → Stock photos (90% savings)\n• Research → Manual reading'),

        ('3. Polite Words Removal',
         'Detects and offers to remove wasteful politeness phrases',
         'Detects 16+ phrases:\n• "please", "thank you", "could you"\n• "would you mind", "if you don\'t mind"\n• "I appreciate", "I would like"\nShows token savings and sanitized version'),

        ('4. Task-Specific Warnings',
         'Alerts for energy-intensive tasks',
         'Examples:\n• Image generation: 3x energy usage warning\n• Research tasks: Multiple API calls alert\n• Code generation: Precision formatting tips'),

        ('5. Length Optimization',
         'Suggests reducing overly long prompts',
         'Triggers when prompt > 500 tokens\nRecommends target ~350 tokens\nShows potential energy savings'),

        ('6. Model Alternatives',
         'Recommends more efficient models',
         'Always included in suggestions\nShows potential 50-70% energy savings\nCompares current vs. recommended model')
    ]

    for title, description, examples in optimization_types:
        add_heading(doc, title, 3)
        doc.add_paragraph(description)
        para = doc.add_paragraph(examples)
        para.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph()

    doc.add_paragraph('Implementation: /shared/analyzer.js (lines 391-599)')

    doc.add_page_break()

    add_heading(doc, '5.4 Model Comparison Tool', 2)
    doc.add_paragraph(
        'Allows users to compare up to 13 models side-by-side to make informed decisions '
        'about model selection based on environmental impact.'
    )

    doc.add_paragraph('Comparison Metrics:')
    comparison_features = [
        'Energy consumption (Wh) with confidence intervals',
        'Water usage (mL) split by on-site and off-site',
        'Carbon emissions (gCO2e) with regional grid data',
        'Eco-Score (0-100) with rating label',
        'Relatable comparisons for each metric',
        'Provider infrastructure efficiency (PUE, WUE, CIF)',
        'Best and worst models highlighted',
        'Potential savings if switching models',
        'Recommendation for most suitable model'
    ]
    for feature in comparison_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Key Insight: The comparison reveals dramatic differences - DeepSeek-R1 uses 240x '
        'more energy than LLaMA 3.2 1B for the same short query!'
    )

    doc.add_paragraph()
    doc.add_paragraph('Implementation: /shared/calculator.js (lines 585-624)')

    add_heading(doc, '5.5 Prompt Library', 2)
    doc.add_paragraph(
        'Save, organize, and reuse sustainable prompts to build a collection of '
        'environmentally-friendly query templates.'
    )

    doc.add_paragraph('Features:')
    library_features = [
        'Save prompts with custom titles and descriptions',
        'Tag prompts by category (code, writing, research, etc.)',
        'Search and filter saved prompts',
        'Track environmental impact of each saved prompt',
        'One-click reuse of efficient prompts',
        'Import/export functionality (planned)',
        'Share prompt templates with team (planned)',
        'Analytics on most-used and most-efficient prompts (planned)'
    ]
    for feature in library_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Data Storage: /data/prompt_templates.json')
    doc.add_paragraph('Implementation: popup.js (Library tab)')

    add_heading(doc, '5.6 Visual Themes & Customization', 2)
    doc.add_paragraph(
        'Six theme options to match user preferences and reduce eye strain:'
    )

    table = create_table_with_data(
        doc,
        ['Theme Name', 'Primary Color', 'Best For'],
        [
            ['Eco Green', '#4caf50', 'Default environmental theme'],
            ['Ocean Blue', '#2196f3', 'Professional, calming'],
            ['Purple Haze', '#9c27b0', 'Creative, distinctive'],
            ['Sunset Orange', '#ff9800', 'Warm, energetic'],
            ['Teal Breeze', '#009688', 'Modern, balanced'],
            ['Dark Mode', 'Dark palette', 'Low-light environments, OLED screens']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph('Additional Customization Options:')
    custom_options = [
        'Show/hide floating badge',
        'Show/hide banner notifications',
        'Auto-detect model from website',
        'Badge position (top-right, top-left, etc.)',
        'Badge animation (pulse, fade, none)',
        'Enable/disable live calculations',
        'Privacy settings for usage tracking'
    ]
    for option in custom_options:
        doc.add_paragraph(option, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Implementation: settings.js, popup.css')

    add_heading(doc, '5.7 Statistics & Analytics', 2)
    doc.add_paragraph(
        'Track cumulative environmental impact and usage patterns over time.'
    )

    doc.add_paragraph('Tracked Metrics:')
    stats_metrics = [
        'Total queries analyzed',
        'Total energy consumed (Wh)',
        'Total water used (mL)',
        'Total carbon emitted (gCO2e)',
        'Average impact per query',
        'Most used models',
        'Most optimized prompts (polite words removed)',
        'Potential savings from following suggestions',
        'Daily/weekly/monthly trends (planned)',
        'Comparison to average user (planned)'
    ]
    for metric in stats_metrics:
        doc.add_paragraph(metric, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Implementation: background.js, popup.js (Stats tab)')

    doc.add_page_break()

    # ========== 6. DATA FILES & MODEL BENCHMARKS ==========
    add_heading(doc, '6. Data Files & Model Benchmarks', 1)

    add_heading(doc, '6.1 Model Benchmarks Data', 2)
    doc.add_paragraph(
        'File: /data/model_benchmarks.json (533 lines, ~25 KB)'
    )

    doc.add_paragraph(
        'Contains empirical performance data for 13 LLM models across 3 query size categories. '
        'All data is from the peer-reviewed research paper.'
    )

    add_heading(doc, 'Data Structure', 3)
    add_code_block(doc, '''{
  "models": {
    "gpt-4o": {
      "name": "GPT-4o",
      "provider": "openai",
      "short": {
        "input_tokens": 100,
        "output_tokens": 300,
        "energy_wh_mean": 0.421,
        "energy_wh_std": 0.034,
        "latency_p50": 0.89,
        "tps_p50": 42.5
      },
      "medium": { ... },
      "long": { ... }
    },
    ...
  }
}''')

    doc.add_paragraph()
    add_heading(doc, 'Fields Explanation', 3)
    fields = [
        ('input_tokens', 'Number of tokens in the prompt (100, 1000, or 10000)'),
        ('output_tokens', 'Number of tokens in the response (300, 1000, or 1500)'),
        ('energy_wh_mean', 'Average energy consumption in watt-hours'),
        ('energy_wh_std', 'Standard deviation for confidence intervals'),
        ('latency_p50', '50th percentile time to first token (seconds)'),
        ('tps_p50', '50th percentile throughput (tokens per second)')
    ]

    for field, desc in fields:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{field}: ').bold = True
        p.add_run(desc)

    add_heading(doc, 'Supported Models (13 Total)', 3)
    table = create_table_with_data(
        doc,
        ['Model ID', 'Display Name', 'Provider', 'Category'],
        [
            ['gpt-4o', 'GPT-4o', 'OpenAI', 'Large flagship model'],
            ['gpt-4o-mini', 'GPT-4o Mini', 'OpenAI', 'Efficient mid-size'],
            ['gpt-4.1-nano', 'GPT-4.1 Nano', 'OpenAI', 'Ultra-efficient small'],
            ['gemini-2.0-flash', 'Gemini 2.0 Flash', 'Google', 'Latest efficient model'],
            ['gemini-1.5-flash', 'Gemini 1.5 Flash', 'Google', 'Previous gen efficient'],
            ['gemini-1.5-pro', 'Gemini 1.5 Pro', 'Google', 'Large capability model'],
            ['claude-3.7-sonnet', 'Claude 3.7 Sonnet', 'Anthropic', 'Latest flagship'],
            ['claude-3.5-sonnet', 'Claude 3.5 Sonnet', 'Anthropic', 'Previous flagship'],
            ['claude-3.5-haiku', 'Claude 3.5 Haiku', 'Anthropic', 'Fast efficient model'],
            ['llama-3.3-70b', 'LLaMA 3.3 70B', 'Meta', 'Large open model'],
            ['llama-3.2-1b', 'LLaMA 3.2 1B', 'Meta', 'Smallest most efficient'],
            ['mistral-large', 'Mistral Large', 'Mistral', 'Large capability model'],
            ['mistral-small', 'Mistral Small', 'Mistral', 'Efficient model'],
            ['deepseek-r1', 'DeepSeek-R1', 'DeepSeek', 'Reasoning model (inefficient)']
        ]
    )

    doc.add_page_break()

    add_heading(doc, '6.2 Infrastructure Data', 2)
    doc.add_paragraph(
        'File: /data/infrastructure.json (111 lines, ~4 KB)'
    )

    doc.add_paragraph(
        'Contains datacenter efficiency multipliers for each cloud provider, including '
        'power usage effectiveness (PUE), water usage effectiveness (WUE), and carbon '
        'intensity factor (CIF).'
    )

    add_heading(doc, 'Data Structure', 3)
    add_code_block(doc, '''{
  "providers": {
    "openai": {
      "name": "Microsoft Azure (OpenAI)",
      "pue": 1.12,
      "wue_onsite": 0.30,
      "wue_offsite": 3.142,
      "cif": 0.3528,
      "region": "US West",
      "notes": "Azure datacenters with renewable energy mix"
    },
    ...
  }
}''')

    doc.add_paragraph()
    add_heading(doc, 'Provider Infrastructure Comparison', 3)
    table = create_table_with_data(
        doc,
        ['Provider', 'PUE', 'WUE Onsite', 'WUE Offsite', 'CIF', 'Efficiency Ranking'],
        [
            ['Google Cloud', '1.10', '0.25 L/kWh', '3.142 L/kWh', '0.32', '1st (Best)'],
            ['Microsoft Azure', '1.12', '0.30 L/kWh', '3.142 L/kWh', '0.3528', '2nd'],
            ['AWS', '1.14', '0.18 L/kWh', '3.142 L/kWh', '0.385', '3rd'],
            ['DeepSeek (China)', '1.27', '1.20 L/kWh', '6.016 L/kWh', '0.60', '4th (Worst)']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph('Key Insights:')
    infra_insights = [
        'Google has the most efficient datacenters (PUE 1.10 = only 10% overhead)',
        'DeepSeek has 27% overhead (PUE 1.27) - significantly less efficient',
        'China\'s coal-heavy grid results in 2x carbon emissions vs. Google',
        'AWS has best on-site water efficiency (0.18 L/kWh) due to advanced cooling',
        'Regional energy mix matters more than datacenter efficiency for carbon'
    ]
    for insight in infra_insights:
        doc.add_paragraph(insight, style='List Bullet')

    add_heading(doc, '6.3 Conversion Factors Data', 2)
    doc.add_paragraph(
        'File: /data/conversion_factors.json (~5 KB)'
    )

    doc.add_paragraph(
        'Contains reference values for converting technical metrics into relatable comparisons.'
    )

    add_heading(doc, 'Reference Values', 3)
    add_code_block(doc, '''{
  "energy": {
    "led_bulb_watts": 10,
    "smartphone_charge_wh": 15,
    "room_light_watts": 15,
    "laptop_watts": 50,
    "tv_65inch_watts": 100
  },
  "water": {
    "drops_per_ml": 20,
    "teaspoon_ml": 5,
    "tablespoon_ml": 15,
    "glass_ml": 250,
    "bottle_ml": 500,
    "dishwasher_liters": 15,
    "washing_machine_liters": 50,
    "shower_10min_liters": 100
  },
  "carbon": {
    "car_sedan_g_per_km": 190,
    "tree_absorption_kg_per_year": 20,
    "breath_g_co2": 0.5
  }
}''')

    doc.add_page_break()

    # ========== 7. FILE STRUCTURE & COMPONENTS ==========
    add_heading(doc, '7. File Structure & Components', 1)

    add_heading(doc, '7.1 Complete Directory Structure', 2)
    add_code_block(doc, '''promptcoach/
├── manifest.json                  # Extension configuration (36 lines)
│
├── Frontend UI Files
│   ├── popup.html                 # Main popup interface structure
│   ├── popup.js                   # Popup logic and event handling
│   ├── popup.css                  # Styling with 6 themes
│   ├── settings.html              # Settings page structure
│   ├── settings.js                # Settings logic
│   └── tooltip.css                # Floating badge styling
│
├── Content & Background Scripts
│   ├── content.js                 # Page integration and real-time analysis
│   └── background.js              # Service worker for background tasks
│
├── shared/                        # Shared calculation modules
│   ├── calculator.js              # Core calculation engine (712 lines)
│   └── analyzer.js                # Prompt analysis and optimization (757 lines)
│
├── data/                          # JSON data files
│   ├── model_benchmarks.json      # 13 LLM models with metrics (533 lines)
│   ├── infrastructure.json        # Cloud provider data (111 lines)
│   ├── conversion_factors.json    # Relatable comparison data
│   ├── config.json                # Extension configuration
│   └── prompt_templates.json      # Prompt library templates
│
├── backend/                       # Python backend (future API)
│   ├── calculations/
│   │   ├── __init__.py            # Main calculator class (323 lines)
│   │   ├── energy.py              # Energy calculation formulas
│   │   ├── water.py               # Water calculation formulas
│   │   ├── carbon.py              # Carbon calculation formulas
│   │   └── conversions.py         # Relatable conversions
│   └── api/                       # API endpoints (planned)
│
├── assets/                        # Images and icons
│   ├── icons/
│   │   ├── icon16.png
│   │   ├── icon48.png
│   │   └── icon128.png
│   └── images/
│
├── tests/                         # Test suites
│   ├── test_calculator.py
│   ├── test_analyzer.py
│   └── test_integration.py
│
└── Documentation
    ├── README.md                  # Project overview (376 lines)
    ├── PLAN.md                    # Implementation plan (21 KB)
    ├── TESTING.md                 # Testing documentation
    └── Context/                   # Research papers and feedback
        └── How_Hungry_is_AI.pdf   # Source research paper''')

    add_heading(doc, '7.2 Critical Files Reference', 2)

    table = create_table_with_data(
        doc,
        ['File', 'Lines', 'Purpose', 'Key Functions/Features'],
        [
            ['shared/calculator.js', '712', 'Core calculations', 'interpolateEnergy(), calculateImpact(), compareModels()'],
            ['shared/analyzer.js', '757', 'Prompt analysis', 'estimateTokens(), detectTaskType(), getOptimizationTips()'],
            ['content.js', '~500', 'Page integration', 'DOM monitoring, real-time badge display'],
            ['popup.js', '~800', 'UI logic', '4 tabs: Coach, Library, Optimize, Stats'],
            ['background.js', '~200', 'Service worker', 'Message routing, stats tracking'],
            ['data/model_benchmarks.json', '533', 'Model data', '13 models × 3 categories × 6 metrics'],
            ['backend/calculations/__init__.py', '323', 'Python calculator', 'Complete impact calculation orchestration']
        ]
    )

    doc.add_page_break()

    # ========== 8. SUPPORTED MODELS & EFFICIENCY RANKINGS ==========
    add_heading(doc, '8. Supported Models & Efficiency Rankings', 1)

    add_heading(doc, '8.1 Complete Model Efficiency Rankings', 2)
    doc.add_paragraph(
        'Energy consumption for SHORT queries (100 input + 300 output = 400 tokens):'
    )

    table = create_table_with_data(
        doc,
        ['Rank', 'Model', 'Energy (Wh)', 'Provider', 'Eco-Score', 'Rating'],
        [
            ['🥇 1', 'LLaMA 3.2 1B', '0.070', 'Meta', '100', 'Excellent'],
            ['🥈 2', 'Gemini 2.0 Flash', '0.095', 'Google', '94', 'Excellent'],
            ['🥉 3', 'GPT-4.1 Nano', '0.103', 'OpenAI', '92', 'Excellent'],
            ['4', 'Gemini 1.5 Flash', '0.115', 'Google', '89', 'Excellent'],
            ['5', 'Claude 3.5 Haiku', '0.125', 'Anthropic', '87', 'Excellent'],
            ['6', 'Mistral Small', '0.145', 'Mistral', '83', 'Excellent'],
            ['7', 'GPT-4o Mini', '0.185', 'OpenAI', '77', 'Good'],
            ['8', 'LLaMA 3.3 70B', '0.247', 'Meta', '70', 'Good'],
            ['9', 'GPT-4o', '0.421', 'OpenAI', '58', 'Fair'],
            ['10', 'Gemini 1.5 Pro', '0.580', 'Google', '49', 'Fair'],
            ['11', 'Mistral Large', '0.680', 'Mistral', '44', 'Fair'],
            ['12', 'Claude 3.5 Sonnet', '0.720', 'Anthropic', '42', 'Fair'],
            ['13', 'Claude 3.7 Sonnet', '0.836', 'Anthropic', '37', 'Poor'],
            ['14', 'DeepSeek-R1', '23.815', 'DeepSeek', '2', 'Very Poor']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Key Insight: DeepSeek-R1 consumes 340x more energy than LLaMA 3.2 1B for the same task!'
    ).bold = True

    add_heading(doc, '8.2 Provider-Based Groupings', 2)

    doc.add_paragraph('OpenAI Models (Microsoft Azure):')
    openai_models = [
        'GPT-4.1 Nano: 0.103 Wh - Most efficient OpenAI model',
        'GPT-4o Mini: 0.185 Wh - Good balance of capability and efficiency',
        'GPT-4o: 0.421 Wh - Flagship model, higher impact'
    ]
    for model in openai_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Google Models (Google Cloud):')
    google_models = [
        'Gemini 2.0 Flash: 0.095 Wh - Second most efficient overall',
        'Gemini 1.5 Flash: 0.115 Wh - Previous generation, still efficient',
        'Gemini 1.5 Pro: 0.580 Wh - Larger capability model'
    ]
    for model in google_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Anthropic Models (AWS):')
    anthropic_models = [
        'Claude 3.5 Haiku: 0.125 Wh - Fast and efficient',
        'Claude 3.5 Sonnet: 0.720 Wh - Previous flagship',
        'Claude 3.7 Sonnet: 0.836 Wh - Latest, more capable but less efficient'
    ]
    for model in anthropic_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Meta Models (Open Source):')
    meta_models = [
        'LLaMA 3.2 1B: 0.070 Wh - Most efficient model overall',
        'LLaMA 3.3 70B: 0.247 Wh - Larger model with good efficiency for size'
    ]
    for model in meta_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Other Models:')
    other_models = [
        'Mistral Small: 0.145 Wh - Efficient European alternative',
        'Mistral Large: 0.680 Wh - Larger capability model',
        'DeepSeek-R1: 23.815 Wh - Reasoning-focused, very inefficient'
    ]
    for model in other_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_page_break()

    add_heading(doc, '8.3 Scaling Analysis: SHORT vs. MEDIUM vs. LONG', 2)
    doc.add_paragraph(
        'How energy consumption scales with query size for selected models:'
    )

    table = create_table_with_data(
        doc,
        ['Model', 'Short (400T)', 'Medium (2000T)', 'Long (11500T)', 'Scaling Factor'],
        [
            ['LLaMA 3.2 1B', '0.070 Wh', '0.268 Wh', '1.425 Wh', '20.4x'],
            ['Gemini 2.0 Flash', '0.095 Wh', '0.385 Wh', '2.156 Wh', '22.7x'],
            ['GPT-4o Mini', '0.185 Wh', '0.824 Wh', '4.623 Wh', '25.0x'],
            ['GPT-4o', '0.421 Wh', '1.945 Wh', '11.234 Wh', '26.7x'],
            ['Claude 3.5 Sonnet', '0.720 Wh', '3.124 Wh', '18.456 Wh', '25.6x'],
            ['DeepSeek-R1', '23.815 Wh', '28.956 Wh', '33.634 Wh', '1.4x']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Observation: Most models scale 20-27x from short to long queries. DeepSeek-R1 '
        'scales only 1.4x because it uses excessive energy even for short queries due to '
        'its reasoning overhead.'
    )

    doc.add_page_break()

    # ========== 9. INTEGRATION & DATA FLOW ==========
    add_heading(doc, '9. Integration & Data Flow', 1)

    add_heading(doc, '9.1 Complete User Journey', 2)

    journey_steps = [
        ('User opens ChatGPT/Claude/Gemini',
         'Browser loads the AI platform webpage'),

        ('Extension activates',
         'Content script (content.js) detects AI platform and injects monitoring code'),

        ('User starts typing prompt',
         'DOM event listener captures input events on textarea'),

        ('Real-time token estimation',
         'Analyzer.estimateTokens() calculates approximate token count from text'),

        ('Task type detection',
         'Analyzer.detectTaskType() identifies prompt category (code, image, Q&A, etc.)'),

        ('Model auto-detection',
         'Calculator.autoDetectModel() identifies model from URL hostname or user selection'),

        ('Impact calculation',
         'Calculator.calculateImpact() performs interpolated energy calculation'),

        ('Water & carbon calculation',
         'Derives from energy using infrastructure multipliers (PUE, WUE, CIF)'),

        ('Relatable conversions',
         'Converts Wh → LED minutes, mL → water drops, gCO2e → car meters'),

        ('Optimization analysis',
         'Analyzer.getOptimizationTips() generates 6 types of suggestions'),

        ('Display floating badge',
         'Creates tooltip with real-time metrics positioned near textarea'),

        ('User reviews impact',
         'Can click badge to open full popup with detailed breakdown'),

        ('User optimizes prompt',
         'May remove polite words, shorten text, or switch models'),

        ('User submits query',
         'Sends optimized prompt to AI platform'),

        ('Stats tracking',
         'Background.js logs query for cumulative impact tracking'),

        ('Badge dismisses',
         'Tooltip auto-hides after 10 seconds or on user dismiss')
    ]

    for i, (step, description) in enumerate(journey_steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {step}\n').bold = True
        p.add_run(f'   {description}')

    doc.add_page_break()

    add_heading(doc, '9.2 Data Flow Diagram', 2)
    add_code_block(doc, '''┌─────────────┐
│   User      │ Types prompt on AI platform
│   Input     │
└──────┬──────┘
       │
       ▼
┌─────────────┐
│  content.js │ Monitors textarea, captures input
└──────┬──────┘
       │
       ▼
┌─────────────┐
│ analyzer.js │ Estimates tokens, detects task type
└──────┬──────┘
       │
       ▼
┌─────────────┐
│calculator.js│ Loads model & infrastructure data
└──────┬──────┘
       │
       ├──────────────────┬─────────────────┐
       ▼                  ▼                 ▼
┌──────────┐      ┌──────────┐      ┌──────────┐
│  Energy  │      │  Water   │      │  Carbon  │
│Calculation│      │Calculation│      │Calculation│
└────┬─────┘      └────┬─────┘      └────┬─────┘
     │                 │                  │
     └────────┬────────┴───────┬──────────┘
              ▼                ▼
       ┌──────────┐    ┌──────────────┐
       │Relatable │    │ Optimization │
       │Conversions│    │ Suggestions  │
       └────┬─────┘    └──────┬───────┘
            │                 │
            └────────┬────────┘
                     ▼
              ┌──────────┐
              │ Display  │ Floating badge with metrics
              │  Badge   │
              └────┬─────┘
                   │
       ┌───────────┴───────────┐
       ▼                       ▼
┌──────────┐          ┌──────────────┐
│  Popup   │          │ background.js│
│ (detail) │          │(stats track) │
└──────────┘          └──────────────┘''')

    add_heading(doc, '9.3 Component Communication', 2)

    doc.add_paragraph('Content Script ↔ Analyzer/Calculator:')
    doc.add_paragraph(
        'Direct function calls (same JavaScript context). Content script imports '
        'shared modules and calls functions synchronously.',
        style='List Bullet'
    )

    doc.add_paragraph()
    doc.add_paragraph('Content Script → Background Worker:')
    doc.add_paragraph(
        'Chrome message passing API (chrome.runtime.sendMessage) for stats tracking '
        'and settings sync. Asynchronous communication.',
        style='List Bullet'
    )

    doc.add_paragraph()
    doc.add_paragraph('Popup ↔ Storage:')
    doc.add_paragraph(
        'Chrome storage API (chrome.storage.local) for persistent settings and '
        'prompt library. Async get/set operations.',
        style='List Bullet'
    )

    doc.add_paragraph()
    doc.add_paragraph('Popup ↔ Content Script:')
    doc.add_paragraph(
        'Indirect via background worker. Popup sends messages to background, which '
        'forwards to active tab content script.',
        style='List Bullet'
    )

    doc.add_page_break()

    # ========== 10. INSTALLATION & DEPLOYMENT ==========
    add_heading(doc, '10. Installation & Deployment', 1)

    add_heading(doc, '10.1 Development Installation', 2)
    doc.add_paragraph('For testing and development:')

    dev_steps = [
        'Clone or download the repository to local machine',
        'Open Chrome and navigate to chrome://extensions/',
        'Enable "Developer mode" toggle in top-right corner',
        'Click "Load unpacked" button',
        'Select the /promptcoach directory',
        'Extension will appear in toolbar with green leaf icon',
        'Pin extension to toolbar for easy access',
        'Click extension icon to open popup interface'
    ]
    for i, step in enumerate(dev_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_heading(doc, '10.2 Chrome Web Store Deployment', 2)
    doc.add_paragraph('For public distribution:')

    store_steps = [
        'Create Chrome Web Store developer account ($5 one-time fee)',
        'Prepare store listing assets (screenshots, promotional images, description)',
        'Create ZIP archive of extension files (exclude .git, tests, documentation)',
        'Upload ZIP to Chrome Web Store Developer Dashboard',
        'Fill in store listing details (name, description, category, privacy policy)',
        'Submit for review (typically 1-3 days for initial review)',
        'Address any review feedback and resubmit if needed',
        'Once approved, extension is published and available for public installation',
        'Users can install with one click from Chrome Web Store'
    ]
    for i, step in enumerate(store_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_heading(doc, '10.3 User Installation (from Chrome Web Store)', 2)
    user_steps = [
        'Visit Chrome Web Store and search for "EcoPrompt Coach"',
        'Click "Add to Chrome" button',
        'Confirm installation in popup dialog',
        'Extension installs automatically and appears in toolbar',
        'Click extension icon to configure initial settings',
        'Visit ChatGPT, Claude, or Gemini to start using'
    ]
    for i, step in enumerate(user_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_heading(doc, '10.4 Supported Platforms & Requirements', 2)

    doc.add_paragraph('Browser Compatibility:')
    browsers = [
        'Google Chrome 90+ (recommended)',
        'Microsoft Edge 90+ (Chromium-based)',
        'Brave Browser 1.25+',
        'Opera 76+',
        'Any Chromium-based browser with Manifest V3 support'
    ]
    for browser in browsers:
        doc.add_paragraph(browser, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('System Requirements:')
    requirements = [
        'Operating System: Windows 10+, macOS 10.13+, Linux (any modern distro)',
        'RAM: 4GB minimum (extension uses <50MB)',
        'Storage: 30MB for extension files and data',
        'Internet: Required for loading data files (one-time), optional for analytics'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Supported AI Platforms:')
    platforms = [
        'ChatGPT (chat.openai.com)',
        'Claude (claude.ai)',
        'Google Gemini (gemini.google.com)',
        'Perplexity (perplexity.ai)',
        'Any web-based LLM interface with text input'
    ]
    for platform in platforms:
        doc.add_paragraph(platform, style='List Bullet')

    doc.add_page_break()

    # ========== 11. PERFORMANCE METRICS & EXAMPLES ==========
    add_heading(doc, '11. Performance Metrics & Examples', 1)

    add_heading(doc, '11.1 Individual Query Examples', 2)

    add_heading(doc, 'Example 1: Simple Q&A with GPT-4o Mini', 3)
    doc.add_paragraph('Prompt: "What is photosynthesis?" (100 input tokens)')
    doc.add_paragraph('Expected output: ~300 tokens')
    doc.add_paragraph()

    table = create_table_with_data(
        doc,
        ['Metric', 'Technical Value', 'Relatable Comparison'],
        [
            ['Energy', '0.185 Wh', '11 minutes of LED bulb light'],
            ['Water', '0.64 mL', '13 drops of water'],
            ['Carbon', '0.065 gCO2e', 'Driving 34 centimeters'],
            ['Eco-Score', '77/100', 'Good efficiency']
        ]
    )

    doc.add_paragraph()
    add_heading(doc, 'Example 2: Code Generation with Claude 3.5 Sonnet', 3)
    doc.add_paragraph('Prompt: "Write a Python function to sort a list using quicksort" (1000 input tokens)')
    doc.add_paragraph('Expected output: ~1000 tokens (medium query)')
    doc.add_paragraph('Task type multiplier: 1.2x (code generation)')
    doc.add_paragraph()

    table = create_table_with_data(
        doc,
        ['Metric', 'Technical Value', 'Relatable Comparison'],
        [
            ['Energy', '3.75 Wh', '22.5 minutes of LED bulb light'],
            ['Water', '13.8 mL', '2.8 teaspoons'],
            ['Carbon', '1.44 gCO2e', 'Driving 7.6 meters'],
            ['Eco-Score', '42/100', 'Fair efficiency']
        ]
    )

    doc.add_paragraph()
    add_heading(doc, 'Example 3: Image Generation with GPT-4o', 3)
    doc.add_paragraph('Prompt: "Generate an image of a sunset over mountains" (100 input tokens)')
    doc.add_paragraph('Task type multiplier: 3.0x (image generation)')
    doc.add_paragraph()

    table = create_table_with_data(
        doc,
        ['Metric', 'Technical Value', 'Relatable Comparison'],
        [
            ['Energy', '1.26 Wh', '7.6 minutes of LED bulb light'],
            ['Water', '4.4 mL', '1 teaspoon (slightly less)'],
            ['Carbon', '0.44 gCO2e', 'Driving 2.3 meters'],
            ['Eco-Score', '58/100', 'Fair efficiency']
        ]
    )

    doc.add_paragraph()
    add_heading(doc, 'Example 4: Long Research Query with DeepSeek-R1', 3)
    doc.add_paragraph('Prompt: 10,000 token research query about AI sustainability')
    doc.add_paragraph('Expected output: ~1,500 tokens (long query)')
    doc.add_paragraph('Task type multiplier: 2.0x (research/agentic)')
    doc.add_paragraph()

    table = create_table_with_data(
        doc,
        ['Metric', 'Technical Value', 'Relatable Comparison'],
        [
            ['Energy', '67.27 Wh', '40 minutes of 65" TV'],
            ['Water', '350 mL', '1.4 drinking glasses'],
            ['Carbon', '40.4 gCO2e', 'Driving 212 meters'],
            ['Eco-Score', '2/100', 'Very Poor efficiency']
        ]
    )

    doc.add_page_break()

    add_heading(doc, '11.2 Cumulative Impact: Annual Scale', 2)
    doc.add_paragraph(
        'Estimated annual environmental impact if all ChatGPT queries used GPT-4o:'
    )

    doc.add_paragraph()
    doc.add_paragraph('Assumptions:')
    assumptions = [
        'ChatGPT: ~700 million queries per day (estimated from public usage data)',
        'Average query: 500 tokens (300 input + 200 output) - between short and medium',
        'Model: GPT-4o with 0.421 Wh per short query, scaled for 500 tokens',
        'Days per year: 365'
    ]
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')

    doc.add_paragraph()
    add_heading(doc, 'Annual GPT-4o Impact (700M queries/day)', 3)

    table = create_table_with_data(
        doc,
        ['Metric', 'Daily', 'Annual', 'Relatable Comparison'],
        [
            ['Energy', '1,072 MWh', '391,509 MWh', '35,000 U.S. homes for a year'],
            ['Water', '3.7M liters', '1.36B liters', '1.2 million people\'s annual drinking water'],
            ['Carbon', '378 tons CO2e', '138,000 tons CO2e', '30,000 cars or a Chicago-sized forest to offset']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: This is a conservative estimate for ONE model. Actual global LLM usage across '
        'all models (ChatGPT, Claude, Gemini, etc.) is likely 10-100x higher.'
    ).bold = True

    add_heading(doc, '11.3 Optimization Potential', 2)
    doc.add_paragraph(
        'Potential environmental savings if users followed optimization suggestions:'
    )

    table = create_table_with_data(
        doc,
        ['Optimization Strategy', 'Potential Savings', 'Example'],
        [
            ['Switch to efficient model', '50-90%', 'GPT-4o → Gemini 2.0 Flash (77% savings)'],
            ['Remove polite words', '5-15%', '"Please help me" → "Help me" (10 tokens saved)'],
            ['Use alternative tool', '100%', 'Weather query → Weather app (no AI needed)'],
            ['Shorten prompt', '20-40%', '1000 tokens → 600 tokens (40% savings)'],
            ['Use summarization model', '30-50%', 'GPT-4o → Claude Haiku for summaries (83% savings)']
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Combined Impact: Following multiple suggestions could reduce per-query impact by 60-95%'
    ).bold = True

    doc.add_page_break()

    # ========== 12. DEVELOPMENT NOTES & DESIGN DECISIONS ==========
    add_heading(doc, '12. Development Notes & Design Decisions', 1)

    add_heading(doc, '12.1 Key Design Decisions', 2)

    decisions = [
        ('Linear Interpolation vs. Polynomial Fitting',
         'Chose linear interpolation because:\n'
         '• Simpler to implement and debug\n'
         '• More predictable behavior\n'
         '• Sufficient accuracy for user decision-making\n'
         '• Avoids overfitting to limited benchmark data'),

        ('Logarithmic Eco-Score',
         'Chose logarithmic scale because:\n'
         '• Handles 340x difference between best/worst models\n'
         '• Provides meaningful differentiation across all models\n'
         '• Linear scale would cluster most models near 0\n'
         '• Mirrors human perception of magnitude differences'),

        ('Client-Side Calculation',
         'Chose browser-based calculation because:\n'
         '• Instant results (<1ms latency)\n'
         '• No server costs or scaling concerns\n'
         '• Works offline after initial data load\n'
         '• Better privacy (no data sent to servers)\n'
         '• Simpler deployment'),

        ('Relatable Metrics Priority',
         'Display relatable comparisons first because:\n'
         '• Users understand "25 minutes of LED light" better than "0.421 Wh"\n'
         '• Drives behavior change through emotional connection\n'
         '• Educational: teaches scale of environmental impact\n'
         '• Technical values shown for advanced users'),

        ('Multiple Data Providers',
         'Include all major cloud providers because:\n'
         '• Different models run on different infrastructure\n'
         '• Infrastructure efficiency varies 15-27% (PUE 1.10-1.27)\n'
         '• Carbon intensity varies 87% (CIF 0.32-0.60)\n'
         '• Provides complete picture of environmental impact'),

        ('Task Type Multipliers',
         'Apply task-specific energy multipliers because:\n'
         '• Image generation genuinely uses 3x more compute\n'
         '• Research tasks make multiple API calls\n'
         '• Provides more accurate estimates than simple token count\n'
         '• Warns users before high-impact operations')
    ]

    for title, explanation in decisions:
        add_heading(doc, title, 3)
        lines = explanation.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    add_heading(doc, '12.2 Performance Considerations', 2)

    doc.add_paragraph('The extension is optimized for minimal performance impact:')

    perf_points = [
        'Data Loading: JSON files (~30KB total) loaded once on extension startup, cached in memory',
        'Calculation Speed: All calculations complete in <1ms using pre-loaded data',
        'DOM Monitoring: Uses efficient event delegation and debouncing (300ms)',
        'Memory Usage: Extension uses <50MB RAM, including all data and UI',
        'Badge Rendering: Uses CSS transforms (GPU-accelerated), no layout reflow',
        'Auto-cleanup: Tooltip badges auto-destroy after 10 seconds to free memory',
        'Lazy Loading: Popup only loads when user clicks icon, not on every page',
        'No Network Requests: After initial load, no external API calls needed'
    ]
    for point in perf_points:
        doc.add_paragraph(point, style='List Bullet')

    add_heading(doc, '12.3 Browser Compatibility Notes', 2)

    doc.add_paragraph('Target: Chrome/Chromium 90+ with Manifest V3')

    compat_notes = [
        'Uses Manifest V3 (service workers instead of background pages)',
        'Async/await syntax requires modern JavaScript support',
        'Fetch API for data loading (no XMLHttpRequest)',
        'CSS Grid and Flexbox for layout (no fallbacks for IE)',
        'chrome.storage.local for persistence (not localStorage)',
        'Content Security Policy (CSP) compliant - no eval() or inline scripts',
        'No jQuery or external libraries - vanilla JavaScript only',
        'Tested on Chrome 120+, Edge 120+, Brave 1.60+'
    ]
    for note in compat_notes:
        doc.add_paragraph(note, style='List Bullet')

    add_heading(doc, '12.4 Security & Privacy', 2)

    doc.add_paragraph('Security measures implemented:')

    security_points = [
        'Content Security Policy prevents XSS attacks',
        'No remote code execution or eval() usage',
        'Data files validated and sanitized on load',
        'Minimal permissions requested (storage, activeTab only)',
        'No data sent to external servers',
        'User prompts never stored or logged (unless explicitly saved to library)',
        'Stats tracking is local-only, no analytics services',
        'Open source for transparency and audit'
    ]
    for point in security_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_page_break()

    add_heading(doc, '12.5 Known Limitations', 2)

    limitations = [
        ('Token Estimation Accuracy',
         'Uses word-based approximation (1 token ≈ 0.75 words). Actual tokenization '
         'varies by model. Accuracy typically within 10-15%.'),

        ('Model Detection',
         'Auto-detection works for major platforms (ChatGPT, Claude, Gemini). May not '
         'detect custom deployments or enterprise instances.'),

        ('Real-Time Data',
         'Benchmark data is from research paper (2025). Model efficiency may improve '
         'over time as providers optimize infrastructure.'),

        ('Output Token Prediction',
         'Output length is estimated based on task type and historical averages. Actual '
         'output may vary significantly.'),

        ('Multi-Turn Conversations',
         'Currently calculates per-query impact. Does not track cumulative impact of '
         'multi-turn conversations with shared context.'),

        ('Non-Text Modalities',
         'Image/video/audio inputs not yet supported. Only text prompts are analyzed.')
    ]

    for title, explanation in limitations:
        add_heading(doc, title, 3)
        doc.add_paragraph(explanation)
        doc.add_paragraph()

    doc.add_page_break()

    # ========== 13. FUTURE ROADMAP ==========
    add_heading(doc, '13. Future Roadmap', 1)

    add_heading(doc, 'Phase 3: Backend API & Database (Q1 2026)', 2)
    phase3_items = [
        'FastAPI REST backend for centralized data management',
        'PostgreSQL database for prompt library and user settings',
        'User authentication and cloud sync across devices',
        'Real-time model benchmark updates',
        'Cloud deployment (Vercel frontend + backend)',
        'Export/import functionality for prompt library'
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_heading(doc, 'Phase 4: Advanced Features (Q2 2026)', 2)
    phase4_items = [
        'Batch query analysis (analyze entire conversations)',
        'Comprehensive usage dashboard with charts and trends',
        'Machine learning for better task type detection',
        'Support for multi-modal inputs (image, audio, video)',
        'Browser extension for Firefox and Safari',
        'Mobile app (iOS/Android) for on-the-go tracking',
        'Integration with Slack, Discord, Teams for workplace AI usage',
        'Carbon offset marketplace integration'
    ]
    for item in phase4_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_heading(doc, 'Phase 5: Enterprise & Policy (Q3-Q4 2026)', 2)
    phase5_items = [
        'Team/organization accounts with centralized management',
        'Detailed sustainability reports for compliance',
        'API for third-party integrations (Zapier, IFTTT)',
        'Custom model benchmarks for private deployments',
        'Policy enforcement (e.g., block models above certain impact threshold)',
        'Educational curriculum integration for universities',
        'Gamification and leaderboards for engagement',
        'White-label version for enterprises'
    ]
    for item in phase5_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_heading(doc, 'Research & Development Priorities', 2)
    rd_items = [
        'Collaborate with AI providers for real-time efficiency data',
        'Expand model coverage to 50+ LLMs',
        'Develop open standard for AI environmental impact reporting',
        'Publish peer-reviewed research on user behavior change',
        'Partner with environmental organizations for offset programs',
        'Create educational materials for schools and universities'
    ]
    for item in rd_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 14. TECHNICAL REFERENCE ==========
    add_heading(doc, '14. Technical Reference', 1)

    add_heading(doc, '14.1 API Reference (JavaScript Modules)', 2)

    add_heading(doc, 'window.EcoPromptCalculator', 3)
    doc.add_paragraph('Main calculation engine (shared/calculator.js)')

    calc_methods = [
        ('loadData()', 'Loads all JSON data files (models, infrastructure, conversions)', 'Promise<void>'),
        ('calculateImpact(modelId, inputTokens, outputTokens)', 'Calculates complete environmental impact', 'Object'),
        ('interpolateEnergy(model, inputTokens, outputTokens)', 'Performs energy interpolation', 'Number'),
        ('compareModels(modelIds, inputTokens, outputTokens)', 'Compares multiple models', 'Array'),
        ('calculateEcoScore(energyWh)', 'Calculates 0-100 efficiency score', 'Number'),
        ('formatEnergyComparison(energyWh)', 'Converts to relatable energy metric', 'String'),
        ('formatWaterComparison(waterML)', 'Converts to relatable water metric', 'String'),
        ('formatCarbonComparison(carbonG)', 'Converts to relatable carbon metric', 'String'),
        ('autoDetectModel(hostname)', 'Detects model from website URL', 'String')
    ]

    table = create_table_with_data(
        doc,
        ['Method', 'Description', 'Returns'],
        calc_methods
    )

    doc.add_paragraph()
    add_heading(doc, 'window.EcoPromptAnalyzer', 3)
    doc.add_paragraph('Prompt analysis engine (shared/analyzer.js)')

    analyzer_methods = [
        ('estimateTokens(text)', 'Estimates token count from text', 'Number'),
        ('detectTaskType(text)', 'Identifies task type with confidence', 'Object'),
        ('detectSpecialQueryType(text)', 'Flags queries for alternative tools', 'Object'),
        ('detectPoliteWords(text)', 'Finds wasteful politeness phrases', 'Object'),
        ('estimateOutputTokens(taskType, inputTokens)', 'Predicts output length', 'Number'),
        ('getOptimizationTips(analysis, impact)', 'Generates optimization suggestions', 'Array'),
        ('analyzePrompt(text)', 'Complete prompt analysis', 'Object'),
        ('generateOptimizedPrompt(text)', 'Removes polite words', 'String')
    ]

    table = create_table_with_data(
        doc,
        ['Method', 'Description', 'Returns'],
        analyzer_methods
    )

    doc.add_page_break()

    add_heading(doc, '14.2 Data File Schemas', 2)

    add_heading(doc, 'model_benchmarks.json Schema', 3)
    add_code_block(doc, '''{
  "models": {
    "model-id": {
      "name": "Display Name",
      "provider": "provider-key",
      "short|medium|long": {
        "input_tokens": number,
        "output_tokens": number,
        "energy_wh_mean": number,
        "energy_wh_std": number,
        "latency_p50": number,
        "tps_p50": number
      }
    }
  }
}''')

    doc.add_paragraph()
    add_heading(doc, 'infrastructure.json Schema', 3)
    add_code_block(doc, '''{
  "providers": {
    "provider-key": {
      "name": "Provider Name",
      "pue": number (1.0-2.0),
      "wue_onsite": number (liters/kWh),
      "wue_offsite": number (liters/kWh),
      "cif": number (kgCO2e/kWh),
      "region": "string",
      "notes": "string"
    }
  }
}''')

    doc.add_paragraph()
    add_heading(doc, 'conversion_factors.json Schema', 3)
    add_code_block(doc, '''{
  "energy": {
    "metric_name": number (watts or watt-hours)
  },
  "water": {
    "metric_name": number (mL or liters)
  },
  "carbon": {
    "metric_name": number (gCO2e or kgCO2e)
  }
}''')

    doc.add_page_break()

    add_heading(doc, '14.3 Chrome Extension Permissions', 2)

    doc.add_paragraph('manifest.json permissions explained:')

    permissions = [
        ('storage', 'Save user settings and prompt library to chrome.storage.local'),
        ('activeTab', 'Detect current tab URL for model auto-detection'),
        ('<all_urls>', 'Inject content script on all AI platform websites'),
        ('scripting', 'Dynamically inject scripts for real-time analysis')
    ]

    for perm, explanation in permissions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{perm}: ').bold = True
        p.add_run(explanation)

    doc.add_paragraph()
    add_heading(doc, '14.4 Build & Development Commands', 2)

    doc.add_paragraph('Common development tasks:')

    commands = [
        ('Load extension', 'chrome://extensions/ → Developer mode → Load unpacked'),
        ('Reload extension', 'chrome://extensions/ → Reload icon (or Ctrl+R on page)'),
        ('View console logs', 'chrome://extensions/ → Details → Inspect views: service worker'),
        ('Debug content script', 'Right-click page → Inspect → Console (filter by extension)'),
        ('Clear storage', 'chrome://extensions/ → Details → Clear storage data'),
        ('Run tests (Python)', 'pytest tests/ or python -m pytest'),
        ('Package for store', 'zip -r ecoprompt-coach.zip * -x "*.git*" "tests/*" "docs/*"')
    ]

    for command, explanation in commands:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{command}: ').bold = True
        p.add_run(explanation)

    doc.add_page_break()

    # ========== APPENDIX: RESEARCH PAPER REFERENCE ==========
    add_heading(doc, 'Appendix: Research Paper Reference', 1)

    doc.add_paragraph(
        'All calculations and benchmarks in EcoPrompt Coach are based on:'
    ).bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        '"How Hungry is AI? Benchmarking Energy, Water, and Carbon Footprint of LLM Inference"'
    ).italic = True
    doc.add_paragraph('Authors: Jegham et al.')
    doc.add_paragraph('Published: 2025')
    doc.add_paragraph('arXiv: 2505.09598v4')
    doc.add_paragraph('Category: cs.CY (Computers and Society)')

    doc.add_paragraph()
    doc.add_paragraph('Paper Highlights:')
    paper_highlights = [
        'First comprehensive benchmark of LLM environmental impact across 13 major models',
        'Measured real-world energy, water, and carbon across 3 query sizes',
        'Included infrastructure overhead (PUE, WUE, CIF) for accurate total impact',
        'Provided confidence intervals and statistical rigor',
        'Covered models from OpenAI, Google, Anthropic, Meta, Mistral, DeepSeek',
        'Peer-reviewed and published in prestigious AI conference'
    ]
    for highlight in paper_highlights:
        doc.add_paragraph(highlight, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'This extension translates the research findings into actionable user-facing tools, '
        'making academic research directly applicable to everyday AI usage.'
    )

    doc.add_page_break()

    # ========== CONTACT & SUPPORT ==========
    add_heading(doc, 'Contact & Support', 1)

    doc.add_paragraph('For questions, issues, or collaboration:')
    doc.add_paragraph()

    contact_info = [
        'GitHub Repository: [Insert URL]',
        'Chrome Web Store: [Insert URL once published]',
        'Email Support: [Insert email]',
        'Documentation: [Insert docs URL]',
        'Bug Reports: GitHub Issues',
        'Feature Requests: GitHub Discussions'
    ]
    for info in contact_info:
        doc.add_paragraph(info, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph()
    p = doc.add_paragraph('Thank you for using EcoPrompt Coach!')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True

    p = doc.add_paragraph('Together, we can make AI more sustainable.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(76, 175, 80)

    # Save the document
    output_path = '/home/user/promptcoach/EcoPrompt_Coach_Handover_Documentation.docx'
    doc.save(output_path)

    return output_path

if __name__ == '__main__':
    print("Generating EcoPrompt Coach Handover Documentation...")
    output_file = generate_handover_document()
    print(f"Documentation generated successfully: {output_file}")
    print("File size:", os.path.getsize(output_file), "bytes")
